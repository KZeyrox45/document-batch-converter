"""
DOCX document parser using python-docx.

This module provides functionality to parse Word documents and extract
text, tables, images, and metadata into structured JSON format.
"""

import re
import time
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from docbatch.parsers.base import BaseParser
from docbatch.models import (
    DocumentOutput,
    DocumentMetadata,
    DocumentType,
    Section,
    TableData,
    ImageInfo,
)


class DOCXParser(BaseParser):
    """
    Parser for DOCX documents using python-docx.
    
    Extracts:
    - Text content with paragraph and heading structure
    - Tables with headers and rows
    - Image metadata (dimensions)
    - Document metadata (author, title, etc.)
    """
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    
    def parse(self, filepath: str) -> DocumentOutput:
        """
        Parse a DOCX file and return structured output.
        
        Args:
            filepath: Path to the DOCX file.
            
        Returns:
            DocumentOutput containing parsed content.
        """
        self.clear_warnings()
        start_time = time.perf_counter()
        
        filename = Path(filepath).name
        self.logger.info(f"Parsing DOCX: {filename}")
        
        # Extract all components
        metadata = self.extract_metadata(filepath)
        full_text = self.extract_text(filepath)
        tables = self.extract_tables(filepath)
        images = self.extract_images(filepath)
        
        # Detect sections with proper heading levels
        sections = self._detect_sections_from_headings(filepath)
        
        # If no sections from headings, fall back to pattern detection
        if not sections:
            sections = self.detect_sections(full_text)
        
        elapsed = time.perf_counter() - start_time
        
        return DocumentOutput(
            filename=filename,
            metadata=metadata,
            sections=sections,
            tables=tables,
            images=images,
            warnings=self.warnings.copy(),
            conversion_time=elapsed,
        )
    
    def extract_metadata(self, filepath: str) -> DocumentMetadata:
        """Extract metadata from DOCX file."""
        try:
            doc = Document(filepath)
            props = doc.core_properties
            
            return DocumentMetadata(
                file_type=DocumentType.DOCX.value,
                pages=self._count_pages(doc),  # Approximate
                author=props.author,
                title=props.title,
                subject=props.subject,
                creator=props.author,
                created=str(props.created_date) if props.created_date else None,
                modified=str(props.modified_date) if props.modified_date else None,
            )
        except Exception as e:
            self.add_warning("metadata_error", str(e))
            return DocumentMetadata(file_type=DocumentType.DOCX.value)
    
    def extract_text(self, filepath: str) -> str:
        """Extract all text from DOCX file."""
        text_parts = []
        
        try:
            doc = Document(filepath)
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
            
        except Exception as e:
            self.add_warning("file_error", str(e))
            raise
        
        return '\n\n'.join(text_parts)
    
    def extract_tables(self, filepath: str) -> List[TableData]:
        """Extract all tables from DOCX file."""
        tables = []
        
        try:
            doc = Document(filepath)
            
            for table_idx, table in enumerate(doc.tables):
                try:
                    rows_data = []
                    
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            # Clean whitespace
                            cell_text = re.sub(r'\s+', ' ', cell_text)
                            row_data.append(cell_text)
                        
                        if any(cell for cell in row_data):
                            rows_data.append(row_data)
                    
                    if rows_data:
                        headers = rows_data[0] if rows_data else []
                        rows = rows_data[1:] if len(rows_data) > 1 else []
                        
                        tables.append(TableData(
                            index=table_idx,
                            headers=headers,
                            rows=rows,
                        ))
                        
                except Exception as e:
                    self.add_warning(
                        "table_extraction_error",
                        str(e),
                        f"table {table_idx + 1}"
                    )
                    
        except Exception as e:
            self.add_warning("file_error", str(e))
            raise
        
        return tables
    
    def extract_images(self, filepath: str) -> List[ImageInfo]:
        """Extract image information from DOCX file."""
        images = []
        image_index = 0
        
        try:
            doc = Document(filepath)
            
            # Extract images from document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image dimensions if possible
                        image_data = rel.target_part.blob
                        
                        # Basic image info
                        images.append(ImageInfo(
                            index=image_index,
                            format=self._get_image_format(rel.target_ref),
                        ))
                        image_index += 1
                        
                    except Exception as e:
                        self.add_warning(
                            "image_extraction_error",
                            str(e),
                            f"image {image_index}"
                        )
                        
        except Exception as e:
            self.add_warning("file_error", str(e))
            raise
        
        return images
    
    def _detect_sections_from_headings(self, filepath: str) -> List[Section]:
        """
        Detect sections from Word heading styles.
        
        This method uses Word's built-in heading styles to identify
        section boundaries, providing more accurate detection than
        pattern matching alone.
        """
        sections = []
        
        try:
            doc = Document(filepath)
            
            current_section = None
            current_content = []
            
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else ""
                text = para.text.strip()
                
                # Check if this is a heading
                is_heading = False
                heading_level = 0
                
                if style_name.startswith('Heading'):
                    try:
                        heading_level = int(style_name.split()[-1])
                        is_heading = True
                    except ValueError:
                        pass
                elif style_name.startswith('标题'):  # Chinese "Heading"
                    try:
                        heading_level = int(style_name[-1])
                        is_heading = True
                    except ValueError:
                        pass
                
                if is_heading and text:
                    # Save previous section
                    if current_section:
                        current_section.content = '\n'.join(current_content).strip()
                        if current_section.heading or current_section.content:
                            sections.append(current_section)
                    
                    # Start new section
                    current_section = Section(
                        heading=text,
                        content="",
                        level=heading_level,
                    )
                    current_content = []
                else:
                    # Add to current content
                    if text:
                        current_content.append(text)
            
            # Save last section
            if current_section:
                current_section.content = '\n'.join(current_content).strip()
                if current_section.heading or current_section.content:
                    sections.append(current_section)
                    
        except Exception as e:
            self.add_warning("section_detection_error", str(e))
        
        return sections
    
    def _count_pages(self, doc: Document) -> int:
        """
        Estimate page count for DOCX.
        
        Note: python-docx doesn't provide exact page count.
        This is an approximation based on content.
        """
        # Simple estimation: count paragraphs and tables
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        
        # Rough estimate: ~30 paragraphs per page
        estimated_pages = max(1, (para_count // 30) + table_count)
        
        return estimated_pages
    
    def _get_image_format(self, target_ref: str) -> str:
        """Extract image format from target reference."""
        ext = target_ref.lower().split('.')[-1]
        
        format_map = {
            'png': 'PNG',
            'jpg': 'JPEG',
            'jpeg': 'JPEG',
            'gif': 'GIF',
            'bmp': 'BMP',
            'tiff': 'TIFF',
            'wmf': 'WMF',
            'emf': 'EMF',
        }
        
        return format_map.get(ext, ext.upper())
